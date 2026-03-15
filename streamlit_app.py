import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from datetime import date
import io
import requests

st.set_page_config(page_title="Offerta / Proforma Invoice Generator", layout="wide")

# ─────────────────────────────────────────────
# SUPABASE CONFIG
# ─────────────────────────────────────────────
SUPABASE_URL = "https://lztrggttkgvgjouofibd.supabase.co"
SUPABASE_KEY = "sb_publishable_2kCkVA7G9VdPWiIXBGIFPw_O_0gfReQ"
HEADERS = {
    "apikey": SUPABASE_KEY,
    "Authorization": f"Bearer {SUPABASE_KEY}",
    "Content-Type": "application/json",
}

# ─────────────────────────────────────────────
# AUTH
# ─────────────────────────────────────────────
VALID_USERS = {"admin": "rmd2024", "federico": "rmd2024"}

def login():
    st.title("🔐 Login")
    u = st.text_input("Username")
    p = st.text_input("Password", type="password")
    if st.button("Login"):
        if VALID_USERS.get(u) == p:
            st.session_state["logged_in"] = True
            st.session_state["username"] = u
            st.rerun()
        else:
            st.error("Invalid credentials")

if not st.session_state.get("logged_in"):
    login()
    st.stop()

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────
@st.cache_data(ttl=300)
def fetch_contacts():
    r = requests.get(
        f"{SUPABASE_URL}/rest/v1/contacts?select=id,company,full_name,address,zip,city,region,country&order=company",
        headers=HEADERS,
    )
    return r.json() if r.ok else []

@st.cache_data(ttl=300)
def fetch_items():
    r = requests.get(
        f"{SUPABASE_URL}/rest/v1/items?select=id,name,description,unit_price,unit,category&order=category,name",
        headers=HEADERS,
    )
    return r.json() if r.ok else []

@st.cache_data(ttl=300)
def fetch_offerte_numbers():
    r = requests.get(
        f"{SUPABASE_URL}/rest/v1/offerte?select=number&order=id.desc",
        headers=HEADERS,
    )
    if r.ok:
        return [x["number"] for x in r.json() if x.get("number")]
    return []

def save_offerta_number(number: str, client: str):
    fetch_offerte_numbers.clear()
    requests.post(
        f"{SUPABASE_URL}/rest/v1/offerte",
        headers=HEADERS,
        json={"number": number, "client": client},
    )

def format_price_it(value: float) -> str:
    """Italian price format: 1.000,– or 1.000,50"""
    if value == int(value):
        integer = f"{int(value):,}".replace(",", ".")
        return f"{integer},\u2013"
    else:
        s = f"{value:.2f}"
        parts = s.split(".")
        integer = f"{int(parts[0]):,}".replace(",", ".")
        return f"{integer},{parts[1]}"

def suggest_next_number(existing: list) -> str:
    year = str(date.today().year)[2:]
    nums = []
    for n in existing:
        try:
            nums.append(int(n.split("/")[0]))
        except Exception:
            pass
    next_n = max(nums) + 1 if nums else 1
    return f"{next_n:03d}/{year}"

def collapse_para(para):
    """Make a paragraph take zero space."""
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "120")
    spacing.set(qn("w:lineRule"), "exact")
    pPr.append(spacing)
    for run in para.runs:
        run.font.size = Pt(1)

# ─────────────────────────────────────────────
# LANGUAGE SELECTOR
# ─────────────────────────────────────────────
lang = st.sidebar.selectbox("🌐 Language / Lingua", ["Italiano", "English"])
is_eng = lang == "English"

st.sidebar.write(f"👤 {st.session_state['username']}")
if st.sidebar.button("Logout"):
    st.session_state.clear()
    st.rerun()

if is_eng:
    st.title("📄 Proforma Invoice Generator")
    lbl_company      = "Company *"
    lbl_fullname     = "Full Name (optional)"
    lbl_attn         = "Include 'To the attn. of' line?"
    lbl_salutation   = "Salutation"
    lbl_offer_no     = "Offer Number"
    lbl_notes        = "Notes / Comments"
    lbl_payment      = "Payment Terms"
    lbl_delivery_t   = "Delivery Terms (e.g. EXW, FOB)"
    lbl_delivery_d   = "Delivery Time"
    lbl_packing      = "Packing"
    lbl_shipment     = "Shipment"
    lbl_hs           = "HS Code"
    lbl_generate     = "Generate Proforma Invoice"
    lbl_download     = "⬇️ Download Proforma Invoice"
    template_file    = "offerta_template_eng.docx"
    doc_name_prefix  = "ProformaInvoice"
else:
    st.title("📄 Generatore Offerta")
    lbl_company      = "Azienda *"
    lbl_fullname     = "Nome completo (opzionale)"
    lbl_attn         = "Includere riga 'To the attn. of'?"
    lbl_salutation   = "Titolo"
    lbl_offer_no     = "Numero Offerta"
    lbl_notes        = "Note e commenti"
    lbl_payment      = "Condizioni di pagamento"
    lbl_delivery_t   = "Resa (es. EXW, FOB, CIF)"
    lbl_delivery_d   = "Tempi di consegna"
    lbl_packing      = "Imballo"
    lbl_shipment     = "Spedizione"
    lbl_hs           = "Codice HS"
    lbl_generate     = "Genera Offerta"
    lbl_download     = "⬇️ Scarica Offerta"
    template_file    = "offerta_template_ita.docx"
    doc_name_prefix  = "Offerta"

# ─────────────────────────────────────────────
# LOAD DATA
# ─────────────────────────────────────────────
contacts         = fetch_contacts()
items            = fetch_items()
existing_numbers = fetch_offerte_numbers()

# ─────────────────────────────────────────────
# SECTION 1 – CLIENT
# ─────────────────────────────────────────────
st.subheader("1. Cliente / Client")

contact_names = [
    c["company"] + (f" – {c['full_name']}" if c.get("full_name") else "")
    for c in contacts
]

if contact_names:
    sel_idx = st.selectbox(
        "Seleziona contatto / Select contact",
        range(len(contact_names)),
        format_func=lambda i: contact_names[i],
    )
    sel = contacts[sel_idx]
else:
    sel = {}
    st.info("No contacts in database.")

col1, col2 = st.columns(2)
with col1:
    company  = st.text_input(lbl_company,              value=sel.get("company", ""))
    address  = st.text_input("Address / Indirizzo",    value=sel.get("address", ""))
    zip_code = st.text_input("ZIP / CAP",              value=sel.get("zip", ""))
with col2:
    city    = st.text_input("City / Città",            value=sel.get("city", ""))
    region  = st.text_input("Region / Regione",        value=sel.get("region", ""))
    country = st.text_input("Country / Paese",         value=sel.get("country", ""))

doc_date = st.date_input("Data / Date", value=date.today(), format="DD/MM/YYYY")

include_attn = st.checkbox(lbl_attn, value=False)
salutation = ""
full_name  = ""
if include_attn:
    col3, col4 = st.columns([1, 3])
    with col3:
        salutation = st.selectbox(lbl_salutation, ["Mr.", "Ms.", "Dr.", "Prof."])
    with col4:
        full_name = st.text_input(lbl_fullname, value=sel.get("full_name", "") or "")

# ─────────────────────────────────────────────
# SECTION 2 – OFFER NUMBER
# ─────────────────────────────────────────────
st.subheader(f"2. {lbl_offer_no}")

suggested    = suggest_next_number(existing_numbers)
offer_number = st.text_input(f"{lbl_offer_no} (es. 001/26)", value=suggested)

number_ok = True
if offer_number in existing_numbers:
    st.error(f"⛔ Numero già usato / Number already exists: {offer_number}")
    number_ok = False
else:
    try:
        if int(offer_number.split("/")[0]) != int(suggested.split("/")[0]):
            st.warning(f"⚠️ Prossimo progressivo suggerito / Next suggested: {suggested}")
    except Exception:
        pass

# ─────────────────────────────────────────────
# SECTION 3 – PRODUCTS
# ─────────────────────────────────────────────
st.subheader("3. Prodotti / Products")

MAX_ROWS = 15
item_labels = [
    f"{it.get('category', '')}: {it['name']}"
    + (f" – {it['description']}" if it.get("description") else "")
    for it in items
]
item_labels_with_empty = ["— vuoto / empty —"] + item_labels

rows = []
for i in range(MAX_ROWS):
    pos = (i + 1) * 10
    with st.expander(f"Riga / Row {pos}", expanded=(i < 3)):
        c1, c2, c3, c4, c5 = st.columns([4, 1, 2, 1, 2])

        with c1:
            sel_item = st.selectbox(
                "Prodotto", item_labels_with_empty,
                key=f"prod_{i}", label_visibility="collapsed"
            )
        with c2:
            qty = st.number_input(
                "Qty", min_value=0.0, step=1.0,
                key=f"qty_{i}", label_visibility="collapsed"
            )
        with c3:
            if sel_item != "— vuoto / empty —":
                idx = item_labels.index(sel_item)
                default_price = float(items[idx].get("unit_price") or 0)
            else:
                default_price = 0.0
            price = st.number_input(
                "Unit price", min_value=0.0, step=0.01,
                value=default_price, key=f"price_{i}", label_visibility="collapsed"
            )
        with c4:
            unit = st.text_input("Unit", value="ISO", key=f"unit_{i}", label_visibility="collapsed")
        with c5:
            total = qty * price
            st.text_input(
                "Total", value=format_price_it(total) if total else "",
                key=f"tot_{i}", disabled=True, label_visibility="collapsed"
            )

        description = ""
        if sel_item != "— vuoto / empty —":
            idx = item_labels.index(sel_item)
            it  = items[idx]
            description = it["name"]
            if it.get("description"):
                description += f" {it['description']}"

        rows.append({
            "pos":         str(pos),
            "description": description,
            "qty":         qty,
            "price":       price,
            "unit":        unit,
            "total":       total,
        })

grand_total = sum(r["total"] for r in rows if r["description"] and r["qty"] > 0)
st.metric("Grand Total / Totale", format_price_it(grand_total) if grand_total else "0,–")

# ─────────────────────────────────────────────
# SECTION 4 – TERMS
# ─────────────────────────────────────────────
st.subheader("4. Termini / Terms")

col5, col6 = st.columns(2)
with col5:
    notes          = st.text_area(lbl_notes, height=80)
    payment        = st.text_input(lbl_payment,    value="30 days / 30 giorni")
    delivery_terms = st.text_input(lbl_delivery_t, value="EXW Schio")
with col6:
    delivery_time  = st.text_input(lbl_delivery_d, value="")
    packing        = st.text_input(lbl_packing,    value="")
    shipment       = st.text_input(lbl_shipment,   value="")
    hs_code        = st.text_input(lbl_hs,         value="")

# ─────────────────────────────────────────────
# GENERATE DOCUMENT
# ─────────────────────────────────────────────
if not company.strip():
    st.warning("⚠️ Company / Azienda è obbligatoria.")

gen_btn = st.button(lbl_generate, type="primary", disabled=(not company.strip() or not number_ok))

if gen_btn:
    try:
        doc      = Document(f"/mnt/user-data/uploads/{template_file}")
        date_str = doc_date.strftime("%d/%m/%Y")

        # ── Paragraphs ──
        for para in doc.paragraphs:
            txt = para.text

            # Date
            if "[DD/MM/" in txt:
                for run in para.runs:
                    if "[DD/MM/" in run.text:
                        run.text = run.text.replace("[DD/MM/'YY]", date_str)\
                                           .replace("[DD/MM/YYYY]", date_str)
                        run.bold = False

            # Company – BOLD
            if "[COMPANY NAME]" in txt:
                for run in para.runs:
                    if "[COMPANY NAME]" in run.text:
                        run.text = run.text.replace("[COMPANY NAME]", company.upper())
                        run.bold = True

            # Address – NOT bold
            if "[Address]" in txt:
                for run in para.runs:
                    if "[Address]" in run.text:
                        run.text = run.text.replace("[Address]", address)
                        run.bold = False

            # ZIP City Region – NOT bold
            if "[Zip]" in txt:
                for run in para.runs:
                    run.text = (run.text
                                .replace("[Zip]", zip_code)
                                .replace("[City]", city)
                                .replace("[Region]", region))
                    run.bold = False

            # Country – NOT bold
            if "[Country]" in txt:
                for run in para.runs:
                    if "[Country]" in run.text:
                        run.text = run.text.replace("[Country]", country)
                        run.bold = False

            # Attn line
            if "To the attn." in txt:
                if include_attn and (full_name.strip() or salutation):
                    for run in para.runs:
                        run.text = run.text\
                            .replace("[Sal.]", salutation)\
                            .replace("[Full Name]", full_name)
                        run.bold = False
                else:
                    collapse_para(para)

            # Offer number – BOLD
            if "[NNN/YY]" in txt:
                for run in para.runs:
                    if "[NNN/YY]" in run.text:
                        run.text = run.text.replace("[NNN/YY]", offer_number)
                    run.bold = True

            # Notes
            for placeholder in ["Notes and comments (ex. VAT excluded)",
                                 "Note e commenti (ex. IVA esclusa)"]:
                if placeholder in txt:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, notes or "")

        # ── Table 0 – Line items ──
        table0 = doc.tables[0]

        def set_cell_text(cell, text):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.text = ""
            if cell.paragraphs:
                para = cell.paragraphs[0]
                if para.runs:
                    para.runs[0].text = text
                else:
                    para.add_run(text)

        for i, row_data in enumerate(rows):
            trow = table0.rows[i + 1]  # row 0 = header
            cells = trow.cells

            if row_data["description"] and row_data["qty"] > 0:
                set_cell_text(cells[0], row_data["pos"])
                set_cell_text(cells[1], row_data["description"])
                qty_val = int(row_data["qty"]) if row_data["qty"] == int(row_data["qty"]) else row_data["qty"]
                set_cell_text(cells[2], str(qty_val))
                set_cell_text(cells[3], format_price_it(row_data["price"]))
                set_cell_text(cells[4], row_data["unit"])
                set_cell_text(cells[5], format_price_it(row_data["total"]))
            else:
                for c in cells:
                    set_cell_text(c, "")

        # Total row (last row)
        total_row   = table0.rows[-1]
        total_cells = total_row.cells
        for c in total_cells[:4]:
            for para in c.paragraphs:
                for run in para.runs:
                    if "[Delivery terms]" in run.text:
                        run.text = run.text.replace("[Delivery terms]", delivery_terms)
        set_cell_text(total_cells[5], format_price_it(grand_total))

        # ── Table 1 – Terms ──
        table1 = doc.tables[1]

        def replace_in_table1(old, new):
            for trow in table1.rows:
                for cell in trow.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if old in run.text:
                                run.text = run.text.replace(old, new)

        replace_in_table1("[HS code]",       hs_code)
        replace_in_table1("[Payment]",       payment)
        replace_in_table1("[Delivery terms]",delivery_terms)
        replace_in_table1("[Delivery time]", delivery_time)
        replace_in_table1("[Packing]",       packing)
        replace_in_table1("[Shipments]",     shipment)

        # ── Export ──
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        # ── Save number to Supabase ──
        save_offerta_number(offer_number, company)
        fetch_offerte_numbers.clear()

        filename = f"{doc_name_prefix}_{offer_number.replace('/', '-')}_{company.replace(' ', '_')}.docx"
        st.success(f"✅ Documento generato / Document generated: {filename}")
        st.download_button(
            label=lbl_download,
            data=buf,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    except Exception as e:
        st.error(f"Errore / Error: {e}")
        st.exception(e)
